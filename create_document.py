"""create_document.py

Gera o arquivo `parecer_tecnico.docx` com formatação padrão para Parecer Técnico.

Requisitos atendidos:
- TITLE em maiúsculas (saída) com o valor: "PARECER TÉCNICO - CRIMES NAS CONTESTAÇÕES".
- Cabeçalho: "Belém, 11 de dezembro de 2025".
- Assinatura: "José Ivanildo da Costa Navegantes Junior - OAB 23.953".
- Fonte Arial 12.
- Parágrafos justificados.
- Espaçamento 1,5.

Dependência:
    pip install python-docx

Uso:
    python create_document.py
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


TITLE = "PARECER TÉCNICO - CRIMES NAS CONTESTAÇÕES"
HEADER_TEXT = "Belém, 11 de dezembro de 2025"
SIGNATURE_TEXT = "José Ivanildo da Costa Navegantes Junior - OAB 23.953"
OUTPUT_DOCX = "parecer_tecnico.docx"

# Conteúdo completo do parecer
EMENTA = """Crimes contra a honra – calúnia (art. 138 do CP) – Imputação de fatos definidos como crime em peça de contestação – Dois núcleos com maior robustez jurídica: (i) imputação de perseguição (stalking, art. 147-A do CP) e (ii) imputação de injúria grave contra menor (art. 140 do CP, com agravantes conforme o caso) – Fragilidade de sustentar calúnia com base em alegações genéricas de falsidade/manipulação de prova – Medidas de obtenção de dados: distinção entre interceptação e quebra de sigilo de registros; legitimidade e vias adequadas para provocação do Judiciário (MP e autoridade policial, com participação do querelante/advogado como requerente/instigador)."""

RELATORIO = """Submete-se à análise a viabilidade jurídica de propositura de queixa-crime por calúnia, a partir de afirmações lançadas em contestação judicial, nas quais a parte adversa:

sustenta que a conduta do autor se enquadraria no tipo penal de perseguição (stalking);
atribui ao autor a prática de injúria grave dirigida à filha menor da parte ré;
além disso, tece considerações sobre a fragilidade/autenticidade de provas digitais (áudios/prints), insinuando possível manipulação, sem afirmar categoricamente falsificação;
e pretende-se, ainda, avaliar a possibilidade/estratégia de requisição de dados telefônicos/telemáticos (quebra de sigilo) para prova.
É o relatório."""

FUNDAMENTACAO = """1. Delimitação: calúnia e seus requisitos
A calúnia (art. 138 do Código Penal) exige, em síntese:

imputação a alguém;
de fato definido como crime;
falsamente (isto é, imputação não verdadeira);
com dolo de atribuir o crime (ainda que sob forma indireta, desde que inequívoca).
O ponto central é a existência de imputação suficientemente determinada de fato criminoso. Alegações vagas ou ambíguas tendem a ser requalificadas como exercício do direito de defesa (animus defendendi), sobretudo quando inseridas em peça processual.

2. Núcleos com maior sustentação para queixa-crime por calúnia
2.1. Imputação de perseguição (stalking) – art. 147-A do CP
Quando a contestação afirma que a conduta do autor "se enquadra no tipo penal de perseguição (stalking)", tem-se, em tese:

crime imputado: perseguição (art. 147-A, CP);
imputação direcionada a pessoa determinada: o autor;
potencial de tipicidade em calúnia: por atribuição de fato criminoso em linguagem técnica e conclusiva ("se enquadra no tipo penal").
Caso se demonstre que a narrativa não corresponde aos fatos (ou é construída de forma artificial), este bloco se apresenta como primeiro fundamento forte para a calúnia, por conter imputação jurídica relativamente objetiva e identificável.

2.2. Imputação de injúria grave contra a filha menor – art. 140 do CP
A contestação também atribui ao autor a prática de injúria contra a filha menor da parte ré, narrando ofensas específicas como fato certo e consumado, utilizado inclusive como suporte para pedido indenizatório.

Em tese:

crime imputado: injúria (art. 140, CP), com repercussões conforme circunstâncias e eventual incidência de agravantes;
fato delimitado: ofensas dirigidas à menor com conteúdo determinado;
estrutura idônea para calúnia: imputação de crime contra a honra, se a falsidade for demonstrada.
Este constitui o segundo núcleo forte e recomendável como eixo da queixa-crime.

Conclusão parcial: a estratégia mais sólida é estruturar a persecução penal privada de calúnia nesses dois blocos principais:

imputação de stalking;
imputação de injúria contra menor.
3. Fragilidade de sustentar calúnia por "insinuação" de falsificação/fraude processual
A parte adversa afirma, em linhas gerais, que áudios/prints "podem ser falsos", que "não há garantia de que não houve edição", e que provas digitais são "facilmente manipuláveis", buscando desqualificar a força probatória.

Embora isso possa carregar conteúdo depreciativo e insinuar má-fé, a calúnia exige imputação inequívoca de crime. Se a redação permanecer no plano de:

dúvida sobre autenticidade;
impugnação genérica da prova;
crítica à ausência de ata notarial/perícia;
é comum que o Judiciário enquadre como impugnação probatória e exercício de defesa, reduzindo a viabilidade de tipificação por calúnia (risco de rejeição liminar da queixa ou trancamento por ausência de justa causa / atipicidade).

Conclusão parcial: esse trecho é mais útil como contexto de intenção desmoralizante (dolo, animus ofendendi) do que como núcleo autônomo de calúnia por "uso de documento falso" (art. 304, CP) ou "fraude processual" (art. 347, CP), se não houver atribuição direta e determinada.

4. Quebra de sigilo: distinções e legitimidade para requerer
4.1. Interceptação vs. quebra de registros/dados
Interceptação de comunicações (Lei 9.296/1996): medida invasiva, voltada à investigação criminal/instrução penal, por ordem judicial e dentro de requisitos estritos. Em regra, não se mostra adequada para crimes contra a honra, que têm pena de detenção.
Quebra de sigilo de dados/registros (registros telefônicos, ERBs, logs, IP, registros de conexão e de acesso a aplicações – Marco Civil da Internet, Lei 12.965/2014): pode ser requerida judicialmente com fundamentação, inclusive para identificação de elementos externos (metadados), a depender do caso.
4.2. Quem pede e quem determina
Quem determina: somente o juiz, por decisão fundamentada.
Quem normalmente requer na esfera penal: Ministério Público e/ou autoridade policial (delegado) conforme a fase (investigação/processo).
Atuação do advogado/querelante: não "determina" a quebra, mas pode provocar:
por requerimentos fundamentados ao delegado (para que represente ao juízo);
por representação ao MP (para que requeira a medida);
e, na ação penal privada, mediante peticionamento ao juízo com fundamentação e pedido de manifestação do MP como fiscal da lei, conforme a prática do foro.
4.3. Estratégia recomendada
Ajuizar a queixa-crime por calúnia centrada nos dois núcleos mais robustos (stalking e injúria contra menor).
Paralelamente, formalizar representação ao Ministério Público, instruída com os autos/documentos, demonstrando a necessidade de dados e a impossibilidade de obtenção direta pelo particular, requerendo que o MP avalie e, se cabível, requeira ao juízo as diligências de obtenção de registros.
Havendo boletim/inquérito, peticionar também nos autos do procedimento policial para requerer diligências."""

CONCLUSAO = """Diante do exposto, opina-se:

Pela viabilidade jurídica de queixa-crime por calúnia com maior chance de sustentação quando fundada em dois eixos:
a) imputação de crime de perseguição (stalking, art. 147-A do CP);
b) imputação de injúria atribuída como fato certo contra a filha menor (art. 140 do CP), desde que demonstrável a falsidade/deturpação relevante do fato narrado.

Pela não recomendação de tratar, como núcleo principal de calúnia, as alegações genéricas de possível manipulação de prova digital, por risco de caracterização como mera impugnação probatória e/ou animus defendendi, admitindo-se seu uso apenas como contexto elucidativo do dolo e da estratégia de desmoralização.

Quanto à prova por dados: interceptação tende a ser inadequada; já a quebra de sigilo de registros/dados pode ser buscada por via judicial, devendo o advogado atuar mediante provocação formal ao MP e/ou autoridade policial, e requerimentos fundamentados ao juízo na medida da pertinência, sem promessa de resultado por depender de decisão judicial."""


def _set_run_font(run) -> None:
    """Define a fonte Arial 12 para o run."""
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(12)


def add_formatted_paragraph(document: Document, text: str) -> None:
    """Adiciona um parágrafo justificado, Arial 12, espaçamento 1,5."""
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    
    run = p.add_run(text)
    _set_run_font(run)


def add_heading_like(document: Document, text: str) -> None:
    """Adiciona um título/seção em destaque simples (negrito) mantendo formatação base."""
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    
    run = p.add_run(text)
    run.bold = True
    _set_run_font(run)


def add_section(document: Document, section_title: str, section_text: str) -> None:
    """Adiciona uma seção com título em negrito e conteúdo formatado."""
    # Adiciona o título da seção em negrito
    add_heading_like(document, section_title)
    
    # Adiciona o conteúdo da seção
    # Divide por linhas em branco para criar parágrafos
    paragraphs = section_text.strip().split('\n\n')
    for para_text in paragraphs:
        # Remove quebras de linha internas e junta em um único parágrafo
        para_text = ' '.join(line.strip() for line in para_text.split('\n') if line.strip())
        if para_text:
            add_formatted_paragraph(document, para_text)


def generate_docx() -> str:
    """Gera o arquivo parecer_tecnico.docx com o conteúdo completo."""
    doc = Document()
    
    # Título em maiúsculas
    add_heading_like(doc, TITLE.upper())
    
    # Cabeçalho
    add_formatted_paragraph(doc, HEADER_TEXT)
    
    # Adiciona linha em branco
    add_formatted_paragraph(doc, "")
    
    # Seção I - EMENTA
    add_section(doc, "I. EMENTA", EMENTA)
    
    # Adiciona linha em branco
    add_formatted_paragraph(doc, "")
    
    # Seção II - RELATÓRIO
    add_section(doc, "II. RELATÓRIO", RELATORIO)
    
    # Adiciona linha em branco
    add_formatted_paragraph(doc, "")
    
    # Seção III - FUNDAMENTAÇÃO
    add_section(doc, "III. FUNDAMENTAÇÃO", FUNDAMENTACAO)
    
    # Adiciona linha em branco
    add_formatted_paragraph(doc, "")
    
    # Seção IV - CONCLUSÃO
    add_section(doc, "IV. CONCLUSÃO", CONCLUSAO)
    
    # Adiciona linhas em branco antes da assinatura
    add_formatted_paragraph(doc, "")
    add_formatted_paragraph(doc, "")
    
    # Assinatura
    add_formatted_paragraph(doc, SIGNATURE_TEXT)
    
    doc.save(OUTPUT_DOCX)
    print(f"Documento gerado com sucesso: {OUTPUT_DOCX}")
    return OUTPUT_DOCX


if __name__ == "__main__":
    generate_docx()
