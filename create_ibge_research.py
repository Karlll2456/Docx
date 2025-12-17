"""create_ibge_research.py

Gera um documento de pesquisa completo sobre o IBGE e o que cai para APM (Agente de Pesquisas e Mapeamento).

Este script cria um documento DOCX formatado com informações detalhadas sobre:
- O que é o IBGE
- História e missão do IBGE
- Funções e atribuições do IBGE
- O cargo de Agente de Pesquisas e Mapeamento (APM)
- Conteúdo programático do concurso para APM
- Disciplinas cobradas na prova
- Dicas de preparação

Dependência:
    pip install python-docx

Uso:
    python create_ibge_research.py
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUTPUT_DOCX = "pesquisa_ibge_apm.docx"


def _set_run_font(run, size=12, bold=False, color=None):
    """Configura a fonte de um run."""
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_title(doc, text):
    """Adiciona um título principal."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _set_run_font(run, size=16, bold=True, color=RGBColor(0, 51, 102))


def add_heading(doc, text, level=1):
    """Adiciona um cabeçalho de seção."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    
    run = p.add_run(text)
    if level == 1:
        _set_run_font(run, size=14, bold=True, color=RGBColor(0, 76, 153))
    else:
        _set_run_font(run, size=12, bold=True, color=RGBColor(0, 102, 204))


def add_paragraph(doc, text, bold=False, justified=True):
    """Adiciona um parágrafo formatado."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justified else WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    
    run = p.add_run(text)
    _set_run_font(run, bold=bold)


def add_bullet_list(doc, items):
    """Adiciona uma lista com marcadores."""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        run = p.add_run(item)
        _set_run_font(run)


def add_numbered_list(doc, items):
    """Adiciona uma lista numerada."""
    for item in items:
        p = doc.add_paragraph(style='List Number')
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        run = p.add_run(item)
        _set_run_font(run)


def generate_ibge_research_document():
    """Gera o documento completo de pesquisa sobre IBGE e APM."""
    doc = Document()
    
    # TÍTULO
    add_title(doc, "PESQUISA COMPLETA: IBGE E CONCURSO PARA APM")
    add_paragraph(doc, "Agente de Pesquisas e Mapeamento", bold=True, justified=False)
    add_paragraph(doc, "")
    
    # SEÇÃO 1: O QUE É O IBGE
    add_heading(doc, "1. O QUE É O IBGE?", level=1)
    
    add_paragraph(doc, "O Instituto Brasileiro de Geografia e Estatística (IBGE) é a principal instituição pública federal dedicada à produção, análise e disseminação de informações estatísticas, geográficas, cartográficas e ambientais sobre o Brasil.")
    
    add_paragraph(doc, "As atividades do IBGE têm impacto direto no planejamento, execução e avaliação de políticas públicas, pesquisas acadêmicas, estratégias de mercado e tomadas de decisão em todo o país.")
    
    # SEÇÃO 2: HISTÓRIA
    add_heading(doc, "2. HISTÓRIA DO IBGE", level=1)
    
    add_bullet_list(doc, [
        "Origem remonta à Diretoria Geral de Estatística (1871)",
        "Instituto Nacional de Estatística criado em 1934",
        "Criação oficial do IBGE em 1936",
        "Em 1938, incorporou o Conselho Brasileiro de Geografia, consolidando-se como IBGE",
        "Primeiro Censo Demográfico realizado em 1940",
        "Desde então, realiza censos e levantamentos que acompanham a evolução da sociedade brasileira"
    ])
    
    # SEÇÃO 3: MISSÃO
    add_heading(doc, "3. MISSÃO DO IBGE", level=1)
    
    add_paragraph(doc, '"Retratar o Brasil com informações necessárias ao conhecimento da sua realidade e ao exercício da cidadania."', bold=True)
    
    add_paragraph(doc, "Esta missão reflete o compromisso do instituto em fornecer dados confiáveis e acessíveis que fundamentem decisões estratégicas em todas as esferas da sociedade.")
    
    # SEÇÃO 4: FUNÇÕES E ATRIBUIÇÕES
    add_heading(doc, "4. FUNÇÕES E ATRIBUIÇÕES DO IBGE", level=1)
    
    add_heading(doc, "4.1. Informações Estatísticas", level=2)
    add_paragraph(doc, "Produção e análise de informações estatísticas sobre:")
    add_bullet_list(doc, [
        "População (Censo Demográfico)",
        "Economia (PIB, índices econômicos)",
        "Saúde e Educação",
        "Trabalho e Renda (PNAD Contínua)",
        "Inflação e Custo de Vida (IPCA)",
        "Meio Ambiente"
    ])
    
    add_heading(doc, "4.2. Informações Geográficas e Cartográficas", level=2)
    add_bullet_list(doc, [
        "Mapeamento completo do território brasileiro",
        "Organização de bases cartográficas",
        "Estudos ambientais e geográficos",
        "Atualização de divisões político-administrativas"
    ])
    
    add_heading(doc, "4.3. Coordenação e Padronização", level=2)
    add_paragraph(doc, "Coordenação e padronização dos sistemas estatísticos e cartográficos nacionais, garantindo uniformidade e qualidade dos dados em todo o país.")
    
    add_heading(doc, "4.4. Disseminação de Informações", level=2)
    add_paragraph(doc, "Divulgação transparente e acessível de todas as informações por meio de:")
    add_bullet_list(doc, [
        "Publicações impressas e digitais",
        "Gráficos e visualizações de dados",
        "Mapas interativos",
        "Bases de dados online",
        "Portal oficial na internet"
    ])
    
    # SEÇÃO 5: PRINCIPAIS CENSOS E PESQUISAS
    add_heading(doc, "5. PRINCIPAIS CENSOS E PESQUISAS DO IBGE", level=1)
    
    add_heading(doc, "Censo Demográfico", level=2)
    add_paragraph(doc, "Realizado a cada 10 anos, é fundamental para conhecer a população brasileira, sua distribuição geográfica, condições de vida e mudanças sociais ao longo do tempo.")
    
    add_heading(doc, "PNAD Contínua (Pesquisa Nacional por Amostra de Domicílios)", level=2)
    add_paragraph(doc, "Pesquisa regular e contínua sobre trabalho, renda, educação e condições de vida da população brasileira.")
    
    add_heading(doc, "IPCA (Índice de Preços ao Consumidor Amplo)", level=2)
    add_paragraph(doc, "Principal índice de inflação do país, acompanha o custo de vida dos brasileiros e serve de referência para a política monetária.")
    
    add_heading(doc, "Censo Agropecuário", level=2)
    add_paragraph(doc, "Levantamento completo sobre o desenvolvimento rural, agricultura, pecuária e estrutura fundiária do Brasil.")
    
    add_heading(doc, "Outras Pesquisas", level=2)
    add_bullet_list(doc, [
        "Pesquisas sobre educação e saúde",
        "Estudos sobre meio ambiente e sustentabilidade",
        "Levantamentos sobre produtividade industrial",
        "Pesquisas municipais e regionais",
        "Estudos econômicos setoriais"
    ])
    
    # SEÇÃO 6: IMPORTÂNCIA SOCIAL E POLÍTICA
    add_heading(doc, "6. IMPORTÂNCIA SOCIAL E POLÍTICA DO IBGE", level=1)
    
    add_bullet_list(doc, [
        "Fundamenta políticas públicas em todas as esferas governamentais",
        "Fornece base para estudos acadêmicos e científicos",
        "Orienta pesquisas de mercado e estratégias empresariais",
        "Subsidia o planejamento territorial e urbano",
        "A colaboração com pesquisas do IBGE é prevista em lei e obrigatória",
        "Coloca o Brasil em posição estratégica nos relatórios globais de desenvolvimento",
        "Garante transparência e democratização do acesso à informação"
    ])
    
    # SEÇÃO 7: O CARGO DE APM
    add_heading(doc, "7. AGENTE DE PESQUISAS E MAPEAMENTO (APM)", level=1)
    
    add_heading(doc, "7.1. O que é o cargo de APM?", level=2)
    add_paragraph(doc, "O Agente de Pesquisas e Mapeamento é um cargo de nível médio no IBGE, fundamental para a produção de estatísticas oficiais e levantamento de informações para políticas públicas. É uma função que exige disposição para trabalho externo e contato direto com a população.")
    
    add_heading(doc, "7.2. Principais Atribuições do APM", level=2)
    
    add_heading(doc, "Coleta de Dados", level=2)
    add_bullet_list(doc, [
        "Visita domicílios, comércios, indústrias e estabelecimentos rurais",
        "Realiza entrevistas presenciais ou por telefone",
        "Preenche questionários impressos ou digitais",
        "Segue protocolos e prazos estabelecidos pelo IBGE"
    ])
    
    add_heading(doc, "Mapeamento e Levantamento Geográfico", level=2)
    add_bullet_list(doc, [
        "Coleta informações cartográficas das áreas de atuação",
        "Registra nomes geográficos e pontos de referência",
        "Identifica características urbanísticas",
        "Colabora com atualização de mapas",
        "Levanta divisões político-administrativas (setores urbanos e rurais)"
    ])
    
    add_heading(doc, "Organização e Transmissão de Dados", level=2)
    add_bullet_list(doc, [
        "Prepara e organiza o material coletado",
        "Transmite ou entrega questionários aos supervisores",
        "Opera sistemas e equipamentos eletrônicos do IBGE",
        "Garante segurança e integridade dos dados coletados"
    ])
    
    add_heading(doc, "Operação de Equipamentos", level=2)
    add_bullet_list(doc, [
        "Opera dispositivos móveis de coleta (tablets e smartphones)",
        "Cuida dos equipamentos cedidos pelo IBGE",
        "Utiliza EPIs quando necessário",
        "Pode dirigir veículos locados pelo instituto",
        "Mantém organização dos materiais de campo"
    ])
    
    add_heading(doc, "7.3. Como é o Dia a Dia do APM?", level=2)
    add_bullet_list(doc, [
        "Trabalho predominantemente de campo",
        "Visitas diárias a locais variados",
        "Deslocamentos entre bairros e localidades",
        "Conversas com moradores e responsáveis por estabelecimentos",
        "Inserção de dados nos sistemas",
        "Soluções para recusa ou ausência de pessoas",
        "Orientação sobre a importância das pesquisas do IBGE",
        "Participação em treinamentos regulares",
        "Atualização sobre mudanças metodológicas"
    ])
    
    add_heading(doc, "7.4. Perfil Ideal para o Cargo", level=2)
    add_bullet_list(doc, [
        "Responsabilidade e disciplina com prazos",
        "Boa organização pessoal",
        "Excelente comunicação interpessoal",
        "Gostar de trabalho externo",
        "Disposição para contato com pessoas",
        "Interesse em questões geográficas e territoriais",
        "Capacidade de trabalhar com autonomia"
    ])
    
    # SEÇÃO 8: CONTEÚDO PROGRAMÁTICO DO CONCURSO
    add_heading(doc, "8. CONTEÚDO PROGRAMÁTICO DO CONCURSO PARA APM", level=1)
    
    add_paragraph(doc, "O concurso do IBGE para Agente de Pesquisas e Mapeamento aborda disciplinas tradicionais de concursos de nível médio, além de conhecimentos específicos importantes para a função.")
    
    # Língua Portuguesa
    add_heading(doc, "8.1. LÍNGUA PORTUGUESA", level=2)
    add_bullet_list(doc, [
        "Interpretação de textos",
        "Compreensão e análise de textos",
        "Tipologia e gêneros textuais",
        "Ortografia oficial",
        "Acentuação gráfica",
        "Pontuação",
        "Classes de palavras",
        "Formação de palavras",
        "Concordância verbal e nominal",
        "Regência verbal e nominal",
        "Colocação pronominal",
        "Crase",
        "Sintaxe da oração e do período",
        "Coesão e coerência textual",
        "Reescrita de frases",
        "Significação de palavras (sinônimos, antônimos, homônimos, parônimos)"
    ])
    
    # Matemática e Raciocínio Lógico
    add_heading(doc, "8.2. MATEMÁTICA E RACIOCÍNIO LÓGICO", level=2)
    
    add_heading(doc, "Matemática:", level=2)
    add_bullet_list(doc, [
        "Razão e proporção",
        "Porcentagem",
        "Regra de três simples e composta",
        "Equações de 1º e 2º grau",
        "Sistemas de equações",
        "Progressões aritméticas e geométricas",
        "Juros simples e compostos",
        "Geometria plana e espacial",
        "Medidas de comprimento, área, volume, massa, tempo",
        "Análise e interpretação de gráficos e tabelas"
    ])
    
    add_heading(doc, "Raciocínio Lógico:", level=2)
    add_bullet_list(doc, [
        "Lógica proposicional",
        "Conectivos lógicos",
        "Tabelas-verdade",
        "Equivalências lógicas",
        "Argumentos lógicos",
        "Sequências numéricas e alfabéticas",
        "Verdades e mentiras",
        "Lógica de primeira ordem",
        "Princípios de contagem",
        "Probabilidade básica"
    ])
    
    # Ética no Serviço Público
    add_heading(doc, "8.3. ÉTICA NO SERVIÇO PÚBLICO", level=2)
    add_bullet_list(doc, [
        "Princípios éticos e morais",
        "Ética e democracia",
        "Exercício da cidadania",
        "Ética e função pública",
        "Código de Ética Profissional do Servidor Público Civil do Poder Executivo Federal",
        "Conduta na administração pública",
        "Legislação relacionada ao funcionalismo público",
        "Deveres e proibições do servidor público",
        "Responsabilidades e penalidades"
    ])
    
    # Noções de Informática
    add_heading(doc, "8.4. NOÇÕES DE INFORMÁTICA", level=2)
    add_bullet_list(doc, [
        "Conceitos básicos de informática",
        "Componentes de hardware e software",
        "Sistemas operacionais (Windows, Linux)",
        "Editores de texto (Microsoft Word, LibreOffice Writer)",
        "Planilhas eletrônicas (Microsoft Excel, LibreOffice Calc)",
        "Apresentações (PowerPoint, Impress)",
        "Conceitos de internet e intranet",
        "Navegadores de internet",
        "Correio eletrônico",
        "Mecanismos de busca",
        "Segurança da informação",
        "Vírus, malware e antivírus",
        "Backup e armazenamento de dados",
        "Computação em nuvem",
        "Ferramentas de colaboração virtual (Google Workspace, Microsoft 365)"
    ])
    
    # Geografia
    add_heading(doc, "8.5. GEOGRAFIA (CONHECIMENTO ESPECÍFICO)", level=2)
    
    add_paragraph(doc, "A Geografia é a disciplina de maior peso e importância para o cargo de APM, sendo fundamental para o desempenho das funções. Os temas principais incluem:")
    
    add_heading(doc, "Geografia Física:", level=2)
    add_bullet_list(doc, [
        "Características do território brasileiro",
        "Relevo, clima, hidrografia e vegetação do Brasil",
        "Biomas brasileiros",
        "Recursos naturais",
        "Questões ambientais",
        "Sustentabilidade e preservação ambiental"
    ])
    
    add_heading(doc, "Cartografia:", level=2)
    add_bullet_list(doc, [
        "Noções básicas de cartografia",
        "Leitura e interpretação de mapas",
        "Escalas cartográficas",
        "Coordenadas geográficas",
        "Projeções cartográficas",
        "Orientação e localização espacial",
        "Fusos horários",
        "Uso de GPS e tecnologias de geolocalização"
    ])
    
    add_heading(doc, "Geografia Humana e Econômica:", level=2)
    add_bullet_list(doc, [
        "Dinâmica populacional brasileira",
        "Crescimento demográfico",
        "Migrações internas e externas",
        "Distribuição da população no território",
        "Urbanização e problemas urbanos",
        "Estrutura econômica do Brasil",
        "Setores da economia (primário, secundário, terciário)",
        "Atividades econômicas regionais",
        "Divisão regional do Brasil",
        "Aspectos socioeconômicos das regiões"
    ])
    
    add_heading(doc, "Geografia Política e Administrativa:", level=2)
    add_bullet_list(doc, [
        "Organização político-administrativa do Brasil",
        "Divisão política: estados e municípios",
        "Regiões geográficas e geoeconômicas",
        "Fronteiras do Brasil",
        "Relações internacionais"
    ])
    
    add_heading(doc, "Temas relacionados à atuação do IBGE:", level=2)
    add_bullet_list(doc, [
        "Coleta e análise de dados estatísticos",
        "Metodologias de pesquisa do IBGE",
        "Importância dos censos",
        "Uso de dados geográficos e estatísticos",
        "Setores censitários",
        "Malha territorial brasileira"
    ])
    
    # SEÇÃO 9: DICAS DE PREPARAÇÃO
    add_heading(doc, "9. DICAS DE PREPARAÇÃO PARA O CONCURSO", level=1)
    
    add_numbered_list(doc, [
        "Foque na Geografia - É a disciplina com maior peso e mais relevante para a função de APM",
        "Estude o edital completo - Conheça todos os tópicos que serão cobrados",
        "Resolva provas anteriores - Especialmente do IBGE e da banca organizadora (geralmente FGV ou IBFC)",
        "Pratique Raciocínio Lógico - É uma matéria que exige bastante prática de questões",
        "Treine informática na prática - Use os programas mencionados no edital",
        "Mantenha-se atualizado - Acompanhe notícias sobre o IBGE e suas pesquisas",
        "Organize um cronograma - Distribua o tempo de estudo entre todas as disciplinas",
        "Faça simulados - Teste seus conhecimentos em condições semelhantes à prova",
        "Revise constantemente - Use técnicas de revisão espaçada",
        "Estude Língua Portuguesa diariamente - É fundamental para todas as outras matérias",
        "Conheça a legislação - Principalmente o Código de Ética do Servidor Público",
        "Pratique leitura de mapas - Fundamental para a função de APM"
    ])
    
    # SEÇÃO 10: ESTRUTURA E BENEFÍCIOS
    add_heading(doc, "10. INFORMAÇÕES SOBRE O CARGO", level=1)
    
    add_heading(doc, "Requisitos:", level=2)
    add_bullet_list(doc, [
        "Ensino Médio completo",
        "Idade mínima de 18 anos",
        "Carteira Nacional de Habilitação (CNH) categoria B (desejável)"
    ])
    
    add_heading(doc, "Remuneração e Benefícios (valores aproximados):", level=2)
    add_bullet_list(doc, [
        "Salário base: aproximadamente R$ 2.676,24 (pode variar conforme edital)",
        "Auxílio-alimentação",
        "Auxílio-transporte",
        "Férias proporcionais",
        "13º salário proporcional"
    ])
    
    add_heading(doc, "Jornada de Trabalho:", level=2)
    add_paragraph(doc, "40 horas semanais, com trabalho predominantemente externo, podendo incluir finais de semana e horários diferenciados conforme as necessidades das pesquisas.")
    
    # SEÇÃO 11: CONCLUSÃO
    add_heading(doc, "11. CONSIDERAÇÕES FINAIS", level=1)
    
    add_paragraph(doc, "O cargo de Agente de Pesquisas e Mapeamento do IBGE é uma excelente oportunidade para quem busca:")
    
    add_bullet_list(doc, [
        "Trabalho dinâmico e variado",
        "Contato direto com a realidade das cidades brasileiras",
        "Contribuir com informações estratégicas para o país",
        "Estabilidade no serviço público",
        "Desenvolvimento profissional na área de pesquisas e estatísticas"
    ])
    
    add_paragraph(doc, "")
    add_paragraph(doc, "A preparação adequada, com foco especial em Geografia e estudo constante de todas as disciplinas, aumenta significativamente as chances de aprovação no concurso.", bold=True)
    
    add_paragraph(doc, "")
    add_paragraph(doc, "Lembre-se: o IBGE é responsável por retratar o Brasil, e você, como APM, será parte fundamental desse processo de conhecimento e desenvolvimento do país!")
    
    add_paragraph(doc, "")
    add_paragraph(doc, "")
    add_paragraph(doc, "___________________________________________")
    add_paragraph(doc, "Documento gerado automaticamente", justified=False)
    add_paragraph(doc, f"Pesquisa sobre IBGE e Concurso para APM", justified=False)
    
    # Salva o documento
    doc.save(OUTPUT_DOCX)
    print(f"✓ Documento '{OUTPUT_DOCX}' gerado com sucesso!")
    return OUTPUT_DOCX


if __name__ == "__main__":
    generate_ibge_research_document()
